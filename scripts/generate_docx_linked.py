#!/usr/bin/env python3
"""Generate formatted RTL Hebrew docx with clickable links from index to full sources."""

import json
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy


# Global bookmark ID counter (must be unique across the document)
_bookmark_id = [0]


def next_bookmark_id():
    _bookmark_id[0] += 1
    return _bookmark_id[0]


def make_bookmark_name(chapter_name, entry_number, prefix="src"):
    """Create a valid bookmark name from chapter + entry number."""
    # Bookmark names must be alphanumeric/underscores, max 40 chars in Word
    # Use a hash-like approach to avoid Hebrew in bookmark names
    ch_hash = abs(hash(chapter_name)) % 100000
    entry_hash = abs(hash(entry_number)) % 100000
    return f"{prefix}_{ch_hash}_{entry_hash}"


def add_bookmark(paragraph, bookmark_name):
    """Add a bookmark (anchor) to a paragraph."""
    bid = next_bookmark_id()
    bookmark_start = parse_xml(
        f'<w:bookmarkStart {nsdecls("w")} w:id="{bid}" w:name="{bookmark_name}"/>'
    )
    bookmark_end = parse_xml(
        f'<w:bookmarkEnd {nsdecls("w")} w:id="{bid}"/>'
    )
    paragraph._element.append(bookmark_start)
    paragraph._element.append(bookmark_end)


def add_hyperlink_to_bookmark(paragraph, bookmark_name, text, font_name='David',
                               font_size=10, bold=False, color='0000FF'):
    """Add a hyperlink that jumps to an internal bookmark."""
    # Build the hyperlink XML element
    hyperlink = parse_xml(
        f'<w:hyperlink {nsdecls("w")} w:anchor="{bookmark_name}"/>'
    )

    # Create a run inside the hyperlink
    run_elem = parse_xml(
        f'<w:r {nsdecls("w")}>'
        f'  <w:rPr>'
        f'    <w:rStyle w:val="Hyperlink"/>'
        f'    <w:rFonts w:cs="{font_name}" w:ascii="{font_name}" w:hAnsi="{font_name}"/>'
        f'    <w:sz w:val="{font_size * 2}"/>'
        f'    <w:szCs w:val="{font_size * 2}"/>'
        f'    <w:color w:val="{color}"/>'
        f'    <w:u w:val="single"/>'
        f'    <w:rtl/>'
        f'  </w:rPr>'
        f'  <w:t xml:space="preserve">{text}</w:t>'
        f'</w:r>'
    )
    if bold:
        rPr = run_elem.find(qn('w:rPr'))
        rPr.append(parse_xml(f'<w:b {nsdecls("w")}/>'))
        rPr.append(parse_xml(f'<w:bCs {nsdecls("w")}/>'))

    hyperlink.append(run_elem)
    paragraph._element.append(hyperlink)


def set_paragraph_rtl(paragraph):
    """Set a paragraph to RTL direction."""
    pPr = paragraph._element.get_or_add_pPr()
    bidi = parse_xml(f'<w:bidi {nsdecls("w")} val="1"/>')
    pPr.append(bidi)


def add_cell_text(cell, text, font_name='David', font_size=10, bold=False,
                  alignment=WD_ALIGN_PARAGRAPH.RIGHT):
    """Add text to a cell with formatting."""
    for p in cell.paragraphs:
        p.clear()

    paragraph = cell.paragraphs[0]
    paragraph.alignment = alignment
    set_paragraph_rtl(paragraph)

    run = paragraph.add_run(text)
    run.font.size = Pt(font_size)
    run.font.bold = bold

    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:cs'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)

    rtl = parse_xml(f'<w:rtl {nsdecls("w")}/>')
    rPr.append(rtl)

    if bold:
        bCs = parse_xml(f'<w:bCs {nsdecls("w")}/>')
        rPr.append(bCs)


def add_cell_hyperlink(cell, bookmark_name, text, font_name='David', font_size=10,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER):
    """Replace cell content with a hyperlink to a bookmark."""
    for p in cell.paragraphs:
        p.clear()

    paragraph = cell.paragraphs[0]
    paragraph.alignment = alignment
    set_paragraph_rtl(paragraph)

    add_hyperlink_to_bookmark(paragraph, bookmark_name, text,
                              font_name=font_name, font_size=font_size)


def set_table_rtl(table):
    """Set table to RTL (bidiVisual)."""
    tblPr = table._element.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        table._element.insert(0, tblPr)
    bidi_visual = parse_xml(f'<w:bidiVisual {nsdecls("w")}/>')
    tblPr.append(bidi_visual)


def set_column_widths(table, widths_cm):
    """Set column widths for a table."""
    tblGrid = table._element.find(qn('w:tblGrid'))
    if tblGrid is not None:
        for col, width in zip(tblGrid.findall(qn('w:gridCol')), widths_cm):
            col.set(qn('w:w'), str(int(width * 567)))


def add_rtl_run(paragraph, text, font_name='David', font_size=10, bold=False):
    """Add an RTL run to a paragraph."""
    run = paragraph.add_run(text)
    run.font.size = Pt(font_size)
    run.font.bold = bold

    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:cs'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)

    rtl = parse_xml(f'<w:rtl {nsdecls("w")}/>')
    rPr.append(rtl)

    if bold:
        bCs = parse_xml(f'<w:bCs {nsdecls("w")}/>')
        rPr.append(bCs)


def build_source_lookup(book_data):
    """Build a lookup: (chapter_name, entry_number) -> list of source texts."""
    lookup = {}
    for chapter in book_data['chapters']:
        ch_name = chapter['name']
        for section in chapter['sections']:
            for entry in section['entries']:
                key = (ch_name, entry['number'])
                texts = []
                for src in entry.get('sources', []):
                    t = src.get('text', '').strip()
                    if t:
                        texts.append(t)
                lookup[key] = {
                    'title': entry.get('title', ''),
                    'texts': texts,
                }
    return lookup


def generate_docx(index_data, book_data, output_path, font_name='David', font_size=10):
    """Generate formatted docx with linked sources."""
    doc = Document()

    # Ensure Hyperlink style exists
    try:
        doc.styles['Hyperlink']
    except KeyError:
        # Create a character style for hyperlinks
        from docx.enum.style import WD_STYLE_TYPE
        hl_style = doc.styles.add_style('Hyperlink', WD_STYLE_TYPE.CHARACTER)
        hl_style.font.color.rgb = None  # Will be set per-run
        hl_style.font.underline = True

    style = doc.styles['Normal']
    style.font.size = Pt(font_size)

    source_lookup = build_source_lookup(book_data)

    for chapter in index_data['chapters']:
        ch_name = chapter['name']

        # Chapter heading
        heading = doc.add_heading(level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = heading.add_run(f"{ch_name} — {chapter['title']}")
        run.font.size = Pt(14)
        set_paragraph_rtl(heading)

        # Collect all entries in this chapter for the sources section
        chapter_entries = []

        for section in chapter['sections']:
            if not section['entries']:
                continue

            # Section heading
            if section['section_name']:
                sec_heading = doc.add_heading(level=2)
                sec_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = sec_heading.add_run(section['section_name'])
                run.font.size = Pt(12)
                set_paragraph_rtl(sec_heading)

            # Create table
            num_rows = 1 + len(section['entries'])
            table = doc.add_table(rows=num_rows, cols=4)
            table.style = 'Table Grid'
            set_table_rtl(table)
            set_column_widths(table, [2.5, 5.3, 5.6, 3.1])

            # Header row
            headers = ['האופן', 'תמצית הענין', 'הטעם', 'מקור']
            for j, header_text in enumerate(headers):
                add_cell_text(
                    table.rows[0].cells[j], header_text,
                    font_name=font_name, font_size=font_size,
                    bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                )

            # Data rows
            for i, entry in enumerate(section['entries']):
                row = table.rows[i + 1]
                entry_num = entry['number']

                src_bookmark = make_bookmark_name(ch_name, entry_num, "src")
                tbl_bookmark = make_bookmark_name(ch_name, entry_num, "tbl")

                # Check if we have source text for this entry
                source_key = (ch_name, entry_num)
                has_source = source_key in source_lookup and source_lookup[source_key]['texts']

                # Column 0: Number - make it a hyperlink if source exists
                if has_source:
                    # Add a bookmark on this cell so we can link back
                    cell_para = row.cells[0].paragraphs[0]
                    cell_para.clear()
                    cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    set_paragraph_rtl(cell_para)
                    add_bookmark(cell_para, tbl_bookmark)
                    add_hyperlink_to_bookmark(cell_para, src_bookmark, f"{entry_num}.",
                                              font_name=font_name, font_size=font_size)
                    chapter_entries.append((entry, source_key))
                else:
                    add_cell_text(
                        row.cells[0], f"{entry_num}.",
                        font_name=font_name, font_size=font_size,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER,
                    )

                # Column 1: Summary
                add_cell_text(
                    row.cells[1], entry.get('summary', entry.get('title', '')),
                    font_name=font_name, font_size=font_size,
                )

                # Column 2: Reason
                add_cell_text(
                    row.cells[2], entry.get('reason', ''),
                    font_name=font_name, font_size=font_size,
                )

                # Column 3: Source
                add_cell_text(
                    row.cells[3], entry.get('source', ''),
                    font_name=font_name, font_size=font_size,
                )

            doc.add_paragraph()

        # --- Full sources section for this chapter ---
        if chapter_entries:
            sources_heading = doc.add_heading(level=2)
            sources_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = sources_heading.add_run('מקורות מלאים')
            run.font.size = Pt(13)
            set_paragraph_rtl(sources_heading)

            for entry, source_key in chapter_entries:
                entry_num = entry['number']
                src_data = source_lookup[source_key]
                src_bookmark = make_bookmark_name(ch_name, entry_num, "src")
                tbl_bookmark = make_bookmark_name(ch_name, entry_num, "tbl")

                # Entry title with bookmark
                title_para = doc.add_paragraph()
                title_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                set_paragraph_rtl(title_para)
                add_bookmark(title_para, src_bookmark)
                title_text = f"{entry_num}. {src_data['title']}"
                add_rtl_run(title_para, title_text, font_name=font_name, font_size=11, bold=True)

                # Full source texts
                full_text = '\n'.join(src_data['texts'])
                text_para = doc.add_paragraph()
                text_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                set_paragraph_rtl(text_para)
                add_rtl_run(text_para, full_text, font_name=font_name, font_size=10)

                # "Back to index" link
                back_para = doc.add_paragraph()
                back_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                set_paragraph_rtl(back_para)
                add_hyperlink_to_bookmark(back_para, tbl_bookmark, '↑ חזרה למפתח',
                                          font_name=font_name, font_size=9,
                                          color='0000AA')

                # Small separator
                sep_para = doc.add_paragraph()
                sep_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                sep_run = sep_para.add_run('─' * 30)
                sep_run.font.size = Pt(8)

            # Page break after full sources section (before next chapter)
            doc.add_page_break()

    doc.save(str(output_path))
    print(f"Saved to {output_path}")


def main():
    repo_root = Path(__file__).parent.parent
    index_path = repo_root / 'data' / 'index.json'
    book_path = repo_root / 'data' / 'book_parsed.json'
    output_path = repo_root / 'output' / 'mafteichos_with_sources.docx'

    if not index_path.exists():
        print("Error: data/index.json not found.")
        sys.exit(1)
    if not book_path.exists():
        print("Error: data/book_parsed.json not found.")
        sys.exit(1)

    with open(index_path, 'r', encoding='utf-8') as f:
        index_data = json.load(f)
    with open(book_path, 'r', encoding='utf-8') as f:
        book_data = json.load(f)

    output_path.parent.mkdir(parents=True, exist_ok=True)

    total_entries = sum(
        len(s['entries'])
        for ch in index_data['chapters']
        for s in ch['sections']
    )

    source_lookup = build_source_lookup(book_data)
    matched = sum(
        1
        for ch in index_data['chapters']
        for s in ch['sections']
        for e in s['entries']
        if (ch['name'], e['number']) in source_lookup and source_lookup[(ch['name'], e['number'])]['texts']
    )

    print(f"Index has {total_entries} entries, {matched} have matching source texts.")
    print("Generating linked docx...")

    generate_docx(index_data, book_data, output_path)


if __name__ == '__main__':
    main()
