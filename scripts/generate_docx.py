#!/usr/bin/env python3
"""Generate formatted RTL Hebrew docx from index.json."""

import json
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


def set_cell_rtl(cell):
    """Set a table cell to RTL direction."""
    for paragraph in cell.paragraphs:
        pPr = paragraph._element.get_or_add_pPr()
        bidi = parse_xml(f'<w:bidi {nsdecls("w")} val="1"/>')
        pPr.append(bidi)


def set_cell_font(cell, font_name='David', font_size=10, bold=False):
    """Set font for all runs in a cell."""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(font_size)
            run.font.bold = bold
            # Set complex script font (for Hebrew)
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
                rPr.insert(0, rFonts)
            rFonts.set(qn('w:cs'), font_name)
            rFonts.set(qn('w:ascii'), font_name)
            rFonts.set(qn('w:hAnsi'), font_name)


def set_cell_alignment(cell, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    """Set paragraph alignment for a cell."""
    for paragraph in cell.paragraphs:
        paragraph.alignment = alignment


def add_cell_text(cell, text, font_name='David', font_size=10, bold=False,
                  alignment=WD_ALIGN_PARAGRAPH.RIGHT):
    """Add text to a cell with formatting."""
    # Clear existing paragraphs
    for p in cell.paragraphs:
        p.clear()

    paragraph = cell.paragraphs[0]
    paragraph.alignment = alignment

    # Set RTL
    pPr = paragraph._element.get_or_add_pPr()
    bidi = parse_xml(f'<w:bidi {nsdecls("w")} val="1"/>')
    pPr.append(bidi)

    run = paragraph.add_run(text)
    run.font.size = Pt(font_size)
    run.font.bold = bold

    # Set complex script font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:cs'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)

    # Set RTL on run
    rtl = parse_xml(f'<w:rtl {nsdecls("w")}/>')
    rPr.append(rtl)

    # Bold for complex script
    if bold:
        bCs = parse_xml(f'<w:bCs {nsdecls("w")}/>')
        rPr.append(bCs)


def set_table_rtl(table):
    """Set table to RTL (bidiVisual)."""
    tblPr = table._element.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        table._element.insert(0, tblPr)
    bidi_visual = parse_xml(f'<w:bidiVisual {nsdecls("w")}/>')
    tblPr.append(bidi_visual)


def set_column_widths(table, widths_cm):
    """Set column widths for a table."""
    tblGrid = table._element.find(qn('w:tblGrid'))
    if tblGrid is not None:
        for col, width in zip(tblGrid.findall(qn('w:gridCol')), widths_cm):
            col.set(qn('w:w'), str(int(width * 567)))  # cm to twips


def generate_docx(index_data, output_path, font_name='David', font_size=10):
    """Generate formatted docx from index data."""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.size = Pt(font_size)

    for chapter in index_data['chapters']:
        # Add chapter heading
        heading = doc.add_heading(level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = heading.add_run(f"{chapter['name']} — {chapter['title']}")
        run.font.size = Pt(14)
        # Set RTL on heading
        pPr = heading._element.get_or_add_pPr()
        bidi = parse_xml(f'<w:bidi {nsdecls("w")} val="1"/>')
        pPr.append(bidi)

        for section in chapter['sections']:
            if not section['entries']:
                continue

            # Add section heading if present
            if section['section_name']:
                sec_heading = doc.add_heading(level=2)
                sec_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = sec_heading.add_run(section['section_name'])
                run.font.size = Pt(12)
                pPr = sec_heading._element.get_or_add_pPr()
                bidi = parse_xml(f'<w:bidi {nsdecls("w")} val="1"/>')
                pPr.append(bidi)

            # Create table: header + entries
            num_rows = 1 + len(section['entries'])
            table = doc.add_table(rows=num_rows, cols=4)
            table.style = 'Table Grid'

            # Set RTL
            set_table_rtl(table)

            # Set column widths (proportional to sample: ~15% / ~32% / ~34% / ~19%)
            # Total page width ~16.5cm
            set_column_widths(table, [2.5, 5.3, 5.6, 3.1])

            # Header row
            headers = ['האופן', 'תמצית הענין', 'הטעם', 'מקור']
            for j, header in enumerate(headers):
                add_cell_text(
                    table.rows[0].cells[j], header,
                    font_name=font_name, font_size=font_size,
                    bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                )

            # Data rows
            for i, entry in enumerate(section['entries']):
                row = table.rows[i + 1]

                # Column 0: Number
                add_cell_text(
                    row.cells[0], f"{entry['number']}.",
                    font_name=font_name, font_size=font_size,
                    alignment=WD_ALIGN_PARAGRAPH.CENTER,
                )

                # Column 1: Summary
                add_cell_text(
                    row.cells[1], entry.get('summary', entry.get('title', '')),
                    font_name=font_name, font_size=font_size,
                )

                # Column 2: Reason
                add_cell_text(
                    row.cells[2], entry.get('reason', ''),
                    font_name=font_name, font_size=font_size,
                )

                # Column 3: Source
                add_cell_text(
                    row.cells[3], entry.get('source', ''),
                    font_name=font_name, font_size=font_size,
                )

            # Add spacing after table
            doc.add_paragraph()

    doc.save(str(output_path))
    print(f"Saved to {output_path}")


def main():
    repo_root = Path(__file__).parent.parent
    index_path = repo_root / 'data' / 'index.json'
    output_path = repo_root / 'output' / 'mafteichos.docx'

    if not index_path.exists():
        print("Error: data/index.json not found. Run generate_index.py first.")
        sys.exit(1)

    with open(index_path, 'r', encoding='utf-8') as f:
        index_data = json.load(f)

    output_path.parent.mkdir(parents=True, exist_ok=True)

    total_entries = sum(
        len(s['entries'])
        for ch in index_data['chapters']
        for s in ch['sections']
    )
    print(f"Generating docx with {total_entries} entries...")

    generate_docx(index_data, output_path)


if __name__ == '__main__':
    main()
